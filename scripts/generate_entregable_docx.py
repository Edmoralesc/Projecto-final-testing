#!/usr/bin/env python3
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx not installed", file=sys.stderr)
    sys.exit(2)

root = Path(__file__).resolve().parents[1]
md_path = root / 'docs' / 'entregable2.md'
docx_path = root / 'docs' / '2.Entregable CI_CD testing pipeline.docx'

text = md_path.read_text(encoding='utf-8')

doc = Document()
# Simple title
doc.add_heading('Entregable 2 — CI/CD Testing Pipeline (FastTicket)', level=1)

for raw_line in text.splitlines():
    line = raw_line.rstrip()
    if not line:
        doc.add_paragraph('')
        continue
    # Headings detected by leading 'A.' 'B.' 'C.' or numbered stages
    if line.startswith('A. '):
        doc.add_heading(line, level=2)
        continue
    if line.startswith('B. '):
        doc.add_heading(line, level=2)
        continue
    if line.startswith('C. '):
        doc.add_heading(line, level=2)
        continue
    if line[:2].isdigit() and line[2:3] == ')':
        doc.add_paragraph(line, style='List Number')
        continue
    if line.lstrip().startswith('- '):
        doc.add_paragraph(line.lstrip()[2:], style='List Bullet')
        continue
    if line.startswith('Entregable 2 —'):
        # skip duplicate title from markdown
        continue
    if line.endswith('≤500 palabras'):
        # omit the helper footer line in docx
        continue
    doc.add_paragraph(line)

# Save
doc.save(str(docx_path))

# Optional: print word count for CI visibility
all_text = '\n'.join(p.text for p in doc.paragraphs)
word_count = len([w for w in all_text.split() if w.strip()])
print(f"DOCX written: {docx_path}")
print(f"Word count (approx): {word_count}")